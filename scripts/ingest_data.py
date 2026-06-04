import os
from pathlib import Path
from typing import List
from github import Github
import git
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.docstore.document import Document as LangchainDocument
from dotenv import load_dotenv

load_dotenv()

class DataIngestion:
    def __init__(self):
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        self.documents = []
    
    def ingest_resume(self, file_path: str):
        """Ingest resume (PDF or DOCX)"""
        print(f"Ingesting resume from {file_path}")
        
        if file_path.endswith('.pdf'):
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages])
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported resume format")
        
        doc = LangchainDocument(
            page_content=text,
            metadata={
                "source": "resume",
                "type": "resume",
                "file": file_path
            }
        )
        
        self.documents.append(doc)
        print(f"✓ Resume ingested: {len(text)} characters")
    
    def ingest_github_repos(self, username: str, repo_names: List[str] = None):
        """Ingest GitHub repositories"""
        print(f"Ingesting GitHub repos for {username}")
        
        token = os.getenv("GITHUB_TOKEN")
        if not token:
            print("Warning: GITHUB_TOKEN not set. Rate limits will apply.")
        
        g = Github(token) if token else Github()
        user = g.get_user(username)
        
        if repo_names:
            repos = [user.get_repo(name) for name in repo_names]
        else:
            repos = list(user.get_repos())[:10]
        
        for repo in repos:
            print(f"  Processing: {repo.name}")
            
            # README
            try:
                readme = repo.get_readme()
                content = readme.decoded_content.decode('utf-8')
                doc = LangchainDocument(
                    page_content=content,
                    metadata={
                        "source": "github",
                        "type": "readme",
                        "repo": repo.name,
                        "url": repo.html_url
                    }
                )
                self.documents.append(doc)
            except:
                pass
            
            # Key files
            key_files = ['package.json', 'requirements.txt', 'setup.py', 'docker-compose.yml', 'Dockerfile']
            
            for file_name in key_files:
                try:
                    file_content = repo.get_contents(file_name)
                    content = file_content.decoded_content.decode('utf-8')
                    doc = LangchainDocument(
                        page_content=f"File: {file_name}\n\n{content}",
                        metadata={
                            "source": "github",
                            "type": "config",
                            "repo": repo.name,
                            "file": file_name,
                            "url": repo.html_url
                        }
                    )
                    self.documents.append(doc)
                except:
                    pass
            
            # Recent commits
            commits = list(repo.get_commits()[:20])
            commit_summary = "\n".join([f"- {c.commit.message[:100]}" for c in commits])
            
            doc = LangchainDocument(
                page_content=f"Recent commits for {repo.name}:\n{commit_summary}",
                metadata={
                    "source": "github",
                    "type": "commits",
                    "repo": repo.name,
                    "url": repo.html_url
                }
            )
            self.documents.append(doc)
            
            print(f"  ✓ {repo.name}: {len(self.documents)} total docs")
    
    def create_vectorstore(self):
        """Create and persist vector store"""
        print(f"\nCreating vector store from {len(self.documents)} documents")
        
        split_docs = self.text_splitter.split_documents(self.documents)
        print(f"Split into {len(split_docs)} chunks")
        
        persist_dir = os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")
        
        vectorstore = Chroma.from_documents(
            documents=split_docs,
            embedding=self.embeddings,
            persist_directory=persist_dir
        )
        
        print(f"✓ Vector store created at {persist_dir}")
        return vectorstore

def main():
    print("=== Data Ingestion for AI Persona ===\n")
    
    ingestion = DataIngestion()
    
    resume_path = os.getenv("RESUME_PATH", "./data/resume.pdf")
    if os.path.exists(resume_path):
        ingestion.ingest_resume(resume_path)
    else:
        print(f"Warning: Resume not found at {resume_path}")
    
    github_username = os.getenv("GITHUB_USERNAME")
    github_repos = os.getenv("GITHUB_REPOS", "").split(",")
    
    if github_username:
        ingestion.ingest_github_repos(
            github_username,
            github_repos if github_repos[0] else None
        )
    else:
        print("Warning: GITHUB_USERNAME not set")
    
    if ingestion.documents:
        ingestion.create_vectorstore()
        print("\n✅ Data ingestion complete!")
    else:
        print("\n❌ No documents ingested. Check your configuration.")

if __name__ == "__main__":
    main()
