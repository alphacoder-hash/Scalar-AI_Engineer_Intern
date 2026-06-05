"""
Ingests resume + GitHub repos into ChromaDB for the RAG engine.
Run:  python scripts/ingest_data_groq.py
"""
import os
import sys
import shutil
from pathlib import Path
from typing import List, Dict
from dotenv import load_dotenv

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

load_dotenv()

from github import Github, GithubException
from sentence_transformers import SentenceTransformer
import chromadb

CODE_EXTENSIONS = {".py", ".ts", ".tsx", ".js", ".jsx", ".md", ".yml",
                   ".yaml", ".toml", ".sh", ".json", ".txt", ".html", ".css"}
SKIP_FILES = {"package-lock.json", "yarn.lock", "pnpm-lock.yaml", ".gitignore",
              ".prettierrc", ".eslintrc", ".eslintignore"}
SKIP_DIRS = {"node_modules", ".git", "dist", "build", "__pycache__",
             "venv", ".venv", "coverage", ".next", "static", ".vercel"}
MAX_FILE_BYTES = 40_000

# ── Document + splitter ───────────────────────────────────────────────────────

class Document:
    def __init__(self, page_content: str, metadata: dict):
        self.page_content = page_content
        self.metadata = metadata


def _split(text: str, chunk_size: int = 900, overlap: int = 180) -> List[str]:
    if len(text) <= chunk_size:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            nl = text.rfind("\n", start, end)
            if nl > start + chunk_size // 2:
                end = nl + 1
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


# ── Ingestion ─────────────────────────────────────────────────────────────────

class DataIngestion:
    def __init__(self):
        print("Loading embeddings model...")
        self.encoder = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
        print("Embeddings ready")
        self.documents: List[Document] = []
        self._seen_content: set = set()   # dedup guard

    def _add(self, content: str, metadata: dict):
        key = content[:200]
        if key in self._seen_content:
            return
        self._seen_content.add(key)
        self.documents.append(Document(page_content=content.strip(), metadata=metadata))

    # ── Resume ────────────────────────────────────────────────────────────────

    def ingest_resume(self, file_path: str):
        if not os.path.exists(file_path):
            print(f"Resume not found at {file_path} — using profile summary")
            self._add_profile_summary()
            return
        ext = Path(file_path).suffix.lower()
        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                text = "\n".join(p.extract_text() or "" for p in PdfReader(file_path).pages)
            elif ext == ".docx":
                from docx import Document as Docx
                text = "\n".join(p.text for p in Docx(file_path).paragraphs)
            else:
                print(f"Unsupported resume format: {ext}")
                return
        except Exception as e:
            print(f"Resume read error: {e}")
            self._add_profile_summary()
            return
        self._add(text, {"source": "resume", "type": "resume",
                         "candidate": os.getenv("CANDIDATE_NAME", "Vaibhav Pandey"),
                         "file": file_path})
        # Also always include the structured summary so critical facts are always retrievable
        self._add_profile_summary()
        print(f"Resume ingested: {len(text)} chars")

    def _add_profile_summary(self):
        """Structured ground-truth document — always in the index."""
        self._add(
            """CANDIDATE PROFILE — Vaibhav Pandey (authoritative ground truth)

CONTACT
  Email: Vpandey1707@gmail.com
  LinkedIn: linkedin.com/in/vaibhav-pandey-4532b8290
  GitHub: github.com/alphacoder-hash
  LeetCode: max rating 1713 | CodeChef: 4-star, max rating 1870

EDUCATION
  B.Tech Computer Science, Class of 2027, Vadodara, India

EXPERIENCE
  Centific Premier Hackathon 2.0 — AI Agent Development Contributor (Apr–May 2026, Hyderabad)
  • Built an AI-powered Business Analyst Agent inside an Agentic SDLC platform
  • Ingested PDF, DOCX, PPTX, email, and meeting-transcript inputs
  • Auto-extracted requirements → generated Features, Epics, User Stories
  • Built Human-in-the-Loop (HITL) validation pipelines with confidence scoring
  • Stack: Python, LLM APIs, FastAPI, document parsers

TECHNICAL SKILLS
  Languages   : Python, TypeScript, JavaScript (ES6+), Java, C++
  AI/ML       : Prompt engineering, RAG, LLM-powered apps, AI evaluation workflows,
                HITL validation, scikit-learn, TF-IDF, semantic similarity, vector DBs
  Frameworks  : FastAPI, React.js, Node.js, Express.js, Django, Streamlit
  Infra/Tools : Docker, Git, GitHub Actions, Vercel, ChromaDB, Gradio, Hugging Face
  Core CS     : DSA, OOP, DBMS, OS, Computer Networks

COMPETITIVE PROGRAMMING
  CodeChef 4-star, max 1870 | LeetCode max 1713

HACKATHONS & AWARDS
  • Smart India Hackathon (SIH) 2025 — institute-level qualifier (200+ teams)
  • Grand Finalist — PU Code Hackathon 2.0 and 3.0
  • Centific Premier Hackathon 2.0 — participant (Hyderabad, 2026)
  • Vadodara Hackathon 6.0 — participant

WHY VAIBHAV FOR SCALER AI ENGINEER
  1. Ships production AI systems (IncidentCommander live on Hugging Face Spaces)
  2. Built AI evaluation frameworks — directly relevant to Scaler's AI product needs
  3. HITL pipeline experience from Centific hackathon — exactly what AI teams run
  4. Strong Python + LLM stack: FastAPI, RAG, ChromaDB, Groq, OpenAI
  5. Full-stack capable (TypeScript, React) — can own features end-to-end
  6. Competitive programmer (LeetCode 1713, CodeChef 1870) — strong CS fundamentals
  7. Proven under pressure: ships working demos at hackathons (IncidentCommander, HotelBookingPro)
""",
            {"source": "resume", "type": "profile_summary",
             "candidate": os.getenv("CANDIDATE_NAME", "Vaibhav Pandey"), "file": "profile_summary"}
        )
        print("Profile summary added")

    # ── GitHub ────────────────────────────────────────────────────────────────

    def ingest_github_repos(self, username: str, repo_names: List[str] = None):
        print(f"\nIngesting GitHub repos for {username}")
        token = os.getenv("GITHUB_TOKEN")
        g = Github(token) if token else Github()
        try:
            user = g.get_user(username)
        except GithubException as e:
            print(f"GitHub user error: {e}")
            return

        repos = []
        if repo_names:
            for name in repo_names:
                try:
                    repos.append(user.get_repo(name))
                    print(f"  Queued: {name}")
                except GithubException as e:
                    print(f"  Skipped {name}: {e}")
        else:
            repos = list(user.get_repos())[:12]

        for repo in repos:
            print(f"\n  → {repo.name}")
            self._ingest_repo(repo)

        print(f"\nTotal documents before chunking: {len(self.documents)}")

    def _ingest_repo(self, repo):
        meta = {"source": "github", "repo": repo.name,
                 "url": repo.html_url, "language": repo.language or ""}

        # Repo summary card
        try:
            desc = repo.description or "No description"
            lang = repo.language or "multiple"
            topics = ", ".join(repo.get_topics()) or "none"
            self._add(
                f"Repository: {repo.name}\n"
                f"URL: {repo.html_url}\n"
                f"Description: {desc}\n"
                f"Primary language: {lang}\n"
                f"Topics: {topics}\n"
                f"Stars: {repo.stargazers_count} | Forks: {repo.forks_count}\n"
                f"Created: {repo.created_at.strftime('%Y-%m-%d')} | "
                f"Updated: {repo.updated_at.strftime('%Y-%m-%d')}",
                {**meta, "type": "repo_summary", "file": "meta"}
            )
        except Exception:
            pass

        # README
        try:
            readme = repo.get_readme()
            content = readme.decoded_content.decode("utf-8")
            self._add(content, {**meta, "type": "readme", "file": "README.md"})
            print(f"    README ({len(content)} chars)")
        except Exception:
            pass

        # Key config files
        for fname in ["package.json", "requirements.txt", "pyproject.toml",
                      "Dockerfile", "docker-compose.yml", "Makefile",
                      "tsconfig.json", ".env.example"]:
            try:
                fc = repo.get_contents(fname)
                content = fc.decoded_content.decode("utf-8")
                self._add(f"File: {fname}\n\n{content}",
                          {**meta, "type": "config", "file": fname})
                print(f"    {fname}")
            except Exception:
                pass

        # All source files (recursive, depth 4)
        self._ingest_files(repo, meta, "", depth=0)

        # Commit history + diffs (last 40)
        self._ingest_commits(repo, meta)

    def _ingest_files(self, repo, meta: Dict, path: str, depth: int):
        if depth > 4:
            return
        try:
            items = repo.get_contents(path or "")
            if not isinstance(items, list):
                items = [items]
            for item in items:
                if item.type == "dir":
                    if item.name not in SKIP_DIRS:
                        self._ingest_files(repo, meta, item.path, depth + 1)
                elif item.type == "file":
                    if item.name in SKIP_FILES:
                        continue
                    if Path(item.name).suffix.lower() not in CODE_EXTENSIONS:
                        continue
                    if item.size > MAX_FILE_BYTES:
                        continue
                    try:
                        content = item.decoded_content.decode("utf-8", errors="replace")
                        if len(content.strip()) < 30:
                            continue
                        self._add(f"File: {item.path}\n\n{content}",
                                  {**meta, "type": "source_code", "file": item.path})
                    except Exception:
                        pass
        except Exception:
            pass

    def _ingest_commits(self, repo, meta: Dict):
        """Ingest commit messages + patch for the last 40 commits."""
        try:
            commits = list(repo.get_commits()[:40])
            log_lines = []
            for c in commits:
                msg = c.commit.message.strip().replace("\n", " ")
                date = c.commit.author.date.strftime("%Y-%m-%d")
                author = c.commit.author.name
                sha = c.sha[:8]
                log_lines.append(f"[{date}] {sha} by {author}: {msg}")
            self._add(
                f"Git log for {repo.name}:\n" + "\n".join(log_lines),
                {**meta, "type": "commits", "file": "git_log"}
            )
            print(f"    {len(commits)} commits logged")

            # Diffs for last 10 commits (file-level changes)
            diff_chunks = []
            for c in commits[:10]:
                try:
                    msg = c.commit.message.strip().split("\n")[0][:120]
                    date = c.commit.author.date.strftime("%Y-%m-%d")
                    sha = c.sha[:8]
                    files_changed = []
                    for f in c.files:
                        patch = (f.patch or "")[:800]
                        files_changed.append(
                            f"  {f.filename} (+{f.additions}/-{f.deletions})\n{patch}"
                        )
                    if files_changed:
                        diff_chunks.append(
                            f"Commit {sha} [{date}]: {msg}\n" + "\n".join(files_changed[:5])
                        )
                except Exception:
                    pass
            if diff_chunks:
                self._add(
                    f"Recent commit diffs for {repo.name}:\n\n" + "\n\n---\n\n".join(diff_chunks),
                    {**meta, "type": "commit_diffs", "file": "git_diffs"}
                )
                print(f"    {len(diff_chunks)} commit diffs")
        except Exception as e:
            print(f"    Commits error: {e}")

    # ── Vector store ──────────────────────────────────────────────────────────

    def create_vectorstore(self):
        if not self.documents:
            print("No documents to process!")
            return None

        chunks: List[Document] = []
        for doc in self.documents:
            for text in _split(doc.page_content):
                chunks.append(Document(page_content=text, metadata=doc.metadata))

        print(f"\n{len(self.documents)} docs → {len(chunks)} chunks")

        persist_dir = os.getenv("CHROMA_PERSIST_DIR", "chroma_db")
        if Path(persist_dir).exists():
            shutil.rmtree(persist_dir)
            print(f"Wiped old vector store at {persist_dir}")

        print("Embedding chunks (may take 3-5 min)...")
        client = chromadb.PersistentClient(path=persist_dir)
        collection = client.create_collection(
            name="persona_docs",
            metadata={"hnsw:space": "cosine"}
        )

        batch_size = 64
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i: i + batch_size]
            texts = [c.page_content for c in batch]
            embeddings = self.encoder.encode(texts, show_progress_bar=False).tolist()
            collection.add(
                documents=texts,
                embeddings=embeddings,
                metadatas=[c.metadata for c in batch],
                ids=[f"chunk_{i + j}" for j in range(len(batch))]
            )
            print(f"  Embedded {min(i + batch_size, len(chunks))}/{len(chunks)}")

        print(f"\nVector store saved at {persist_dir} ({len(chunks)} chunks total)")
        return collection


def main():
    print("=" * 55)
    print("Data Ingestion — Sam AI Persona")
    print("=" * 55)

    ing = DataIngestion()

    resume_path = os.getenv("RESUME_PATH", "./data/Vaibhav_Pandey_Intern_Resume.pdf")
    ing.ingest_resume(resume_path)

    username = os.getenv("GITHUB_USERNAME", "alphacoder-hash")
    repo_names = [r.strip() for r in os.getenv("GITHUB_REPOS", "").split(",") if r.strip()]
    ing.ingest_github_repos(username, repo_names or None)

    if ing.documents:
        ing.create_vectorstore()
        print("\n✓ Ingestion complete")
    else:
        print("\n✗ No documents ingested")


if __name__ == "__main__":
    main()
